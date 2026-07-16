"""
server/services/report.py — AI career report generation

Mirrors the real consultant workflow: feed the actual SOP document
(quy_trinh_chot_case.md) and the actual Master Router prompt template
(master_router_prompt.md), plus the Trong nước deep-dive prompt when
the case routes domestic — exactly the "3 file" pattern, not a
condensed rewrite.

Requires:
    ANTHROPIC_API_KEY environment variable
"""

import base64
import copy
import json
import os
import re
import subprocess
import tempfile
import threading
import urllib.error
import urllib.request
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Generator, Optional
from zoneinfo import ZoneInfo

import anthropic
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm, Inches

MODEL      = "claude-opus-4-7"
MAX_TOKENS = 32000
# NOTE: raised from 20000 after the first real V4.2 test (Hà Quyên,
# 2026-07-13) hit exactly 20000 output tokens and got cut off mid-
# sentence in the final [AUDIT NỘI BỘ] section — the V4.2 SOP's
# required structure (15 numbered sections, per-ngành deep-dives x5,
# O*NET + VSCO tables, full roadmap, Quest section, audit appendix) is
# simply longer than the older report format this was tuned for.
# Claude Opus 4.7 supports up to 128,000 output tokens on the standard
# Messages API (per Anthropic's docs) — 20000 was using under a fifth
# of that. 32000 gives real headroom (that test was maybe a few
# hundred tokens short of finishing) without the much larger worst-
# case cost of maxing out at 128K ($25/million output tokens — 128K
# alone could run ~$3.20 if ever fully used; 32K keeps worst case
# closer to ~$0.80 for output).

PROMPTS_DIRNAME = "prompts"
SOP_FILENAME           = "quy_trinh_chot_case_v4_2.md"
MASTER_ROUTER_FILENAME  = "master_router_prompt_v4_2.md"
TRONG_NUOC_FILENAME     = "prompt_2_5_trong_nuoc_v4_2.md"

# Exact match against the live form's dropdown options for
# "Bạn đang quan tâm hướng nào nhất?" — confirmed from the actual form
# (not fuzzy keywords, to avoid silent misrouting if wording is close
# but not identical, e.g. "Chọn ngành đại học ở VN" vs guessed "trong nước").
#
# NOTE: "Học nghề/College" is mapped to Route B (domestic) as a judgment
# call — it's a domestic vocational/college track, but the SOP files may
# not explicitly address vocational pathways the same way as university
# pathways. Revisit if this doesn't read correctly in testing.
ROUTE_BY_DIRECTION = {
    "du học":                  "A",
    "chọn ngành đại học ở vn": "B",
    "học nghề/college":        "B",
    "chưa rõ, đang tìm hiểu":  "C",
}


def _get_prompts_dir() -> Path:
    """
    Find the src/prompts directory. Tries multiple candidate locations
    since the exact relative depth depends on how the Docker image was built.
    """
    here = Path(__file__).resolve()
    candidates = [
        Path("/app/src/prompts"),
        here.parent.parent / "src" / "prompts",
        here.parent.parent.parent / "src" / "prompts",
        Path.cwd() / "src" / "prompts",
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError(f"Prompts directory not found. Tried: {[str(c) for c in candidates]}")


def _read_file(filename: str) -> str:
    path = _get_prompts_dir() / filename
    if not path.exists():
        raise FileNotFoundError(f"Required file not found: {path}")
    return path.read_text(encoding="utf-8")


def _detect_route(direction: str) -> str:
    """
    Exact-match route detection against real form dropdown values.
    Unknown/unexpected values default to "C" (chưa rõ) rather than
    guessing — safer than silently misrouting, and Route C is already
    designed to handle ambiguity per the SOP.
    """
    key = (direction or "").strip().lower()
    return ROUTE_BY_DIRECTION.get(key, "C")


# ============================================================
# Transcript file handling (Pass 2)
#
# Students can attach one or more transcript files (PDF or image) in
# the main survey. Apps Script fetches these from Drive, base64-
# encodes them, and sends them as transcript_files in the
# generate-report-async payload — see mirror_check_response_trigger.gs.
# Claude reads them directly as real document/image content (not OCR'd
# separately), via Anthropic's native document/image content blocks.
# ============================================================

# Anthropic's Messages API uses a different content-block "type" for
# PDFs ("document") vs. images ("image") — this maps each accepted
# mime type to the correct block type. Anything not in this dict is
# rejected (logged, not sent to Claude) rather than guessed at.
TRANSCRIPT_MIME_TO_BLOCK_TYPE = {
    "application/pdf": "document",
    "image/jpeg": "image",
    "image/png": "image",
    "image/gif": "image",
    "image/webp": "image",
}

TRANSCRIPT_MAX_FILES = 5
# ~20MB of actual file bytes, expressed as a base64-character-count
# limit (base64 inflates size by ~4/3) — a generous cap for scanned
# transcripts/report cards, well above what real report-card photos
# or a multi-page PDF should ever need.
TRANSCRIPT_MAX_TOTAL_BASE64_CHARS = 20 * 1024 * 1024 * 4 // 3


def _build_transcript_content_blocks(transcript_files: Optional[list]) -> list:
    """
    Convert the raw transcript_files list (from the API request — each
    item a dict with filename/mime_type/data) into Anthropic API
    content blocks, ready to prepend to the text prompt.

    Applies the same kind of defensive validation as everywhere else
    in this pipeline: unsupported mime types are skipped (not sent to
    Claude, not a hard failure), and both a file-count cap and a
    total-size cap protect against one oversized/malformed upload
    silently blowing up the request. Any single bad file is skipped
    with a log line rather than failing the whole report — a report
    generated from *some* of the transcript is better than no report
    at all.
    """
    if not transcript_files:
        return []

    blocks = []
    total_base64_chars = 0
    files_included = 0

    for f in transcript_files:
        filename = f.get("filename", "unknown")
        mime_type = f.get("mime_type", "")
        data = f.get("data", "")

        if files_included >= TRANSCRIPT_MAX_FILES:
            print(f"WARNING: Transcript file '{filename}' skipped — already at "
                  f"TRANSCRIPT_MAX_FILES ({TRANSCRIPT_MAX_FILES}) limit.")
            continue

        if mime_type not in TRANSCRIPT_MIME_TO_BLOCK_TYPE:
            print(f"WARNING: Transcript file '{filename}' skipped — unsupported "
                  f"mime type '{mime_type}'. Allowed: {list(TRANSCRIPT_MIME_TO_BLOCK_TYPE.keys())}")
            continue

        if not data:
            print(f"WARNING: Transcript file '{filename}' skipped — no data.")
            continue

        if total_base64_chars + len(data) > TRANSCRIPT_MAX_TOTAL_BASE64_CHARS:
            print(f"WARNING: Transcript file '{filename}' skipped — would exceed "
                  f"total transcript size cap ({TRANSCRIPT_MAX_TOTAL_BASE64_CHARS} base64 chars).")
            continue

        block_type = TRANSCRIPT_MIME_TO_BLOCK_TYPE[mime_type]
        blocks.append({
            "type": block_type,
            "source": {
                "type": "base64",
                "media_type": mime_type,
                "data": data,
            },
        })
        total_base64_chars += len(data)
        files_included += 1

    if blocks:
        print(f"=== Attached {len(blocks)} transcript file(s) as document/image content "
              f"({total_base64_chars} base64 chars total) ===")

    return blocks


# Deterministic letter -> meaning mapping for the Mirror Check choice,
# matching the fixed option list in portrait_prompt.md exactly. Spelling
# this out explicitly in the prompt (rather than sending just the bare
# letter and expecting the model to recall/infer its meaning from the
# option list read much earlier in a 150K+ character prompt) is what
# fixes a real error found in testing: a generation for token
# HN-2026-0011 (choice F = "pha giữa A và C") incorrectly wrote "pha
# giữa A và B" in the final report — apparently conflating the
# mismatch_answer field (which portrait she said DIDN'T fit — B, in
# that case) with one of the two portraits actually in her blend.
# Spelling out the meaning here removes the need for the model to
# recall it at all.
CHOICE_LABEL_MAP = {
    "A": "A (bản A giống em nhất)",
    "B": "B (bản B giống em nhất)",
    "C": "C (bản C giống em nhất)",
    "D": "D (pha giữa A và B)",
    "E": "E (pha giữa B và C)",
    "F": "F (pha giữa A và C)",
    "G": "G (không bản nào giống em lắm)",
}


def _build_mirror_check_block(mirror_check: Optional[Dict[str, Any]]) -> str:
    """
    Build the "Mirror Check response" input block required by the
    updated (V4.2) Master Router / Siêu Prompt 2.5 prompts. Per those
    documents, this block must contain:
      - Micro-portrait app đề xuất mạnh nhất (score_matched)
      - Học sinh chọn micro-portrait nào (student_choice)
      - Câu học sinh tick/highlight là đúng nhất
      - Câu học sinh phản hồi không giống mình nếu có
      - Mirror Fit sơ bộ (color / level)

    If mirror_check is None (student hasn't completed Mirror Check yet,
    or this is a legacy /test-report call made before Mirror Check
    existed), returns the explicit fallback line the V4.2 SOP requires:
    the report must say self-confirmation data is missing and treat the
    result as an opening map needing Quest-verification, rather than
    silently omitting the section.
    """
    if not mirror_check:
        return (
            "Mirror Check response: Chưa có dữ liệu self-confirmation; "
            "kết quả cần đọc như bản đồ mở đầu và cần Quest để xác nhận."
        )

    score_matched_raw = (mirror_check.get("score_matched") or "").strip().upper()
    student_choice_raw = (mirror_check.get("student_choice") or "").strip().upper()

    # Use the spelled-out label whenever the letter is recognized;
    # otherwise fall back to the raw value as-is rather than crashing —
    # an unrecognized letter is a real (if unlikely) data issue that
    # should still let the report generate, just without translation.
    score_matched  = CHOICE_LABEL_MAP.get(score_matched_raw, score_matched_raw or "Không rõ")
    student_choice = CHOICE_LABEL_MAP.get(student_choice_raw, student_choice_raw or "Không rõ")

    highlight = mirror_check.get("highlight_answer") or "Không có"
    mismatch  = mirror_check.get("mismatch_answer") or "Không có"
    fit_color = mirror_check.get("mirror_fit_color") or ""
    fit_level = mirror_check.get("mirror_fit_level") or ""
    fit_combined = f"{fit_color} / {fit_level}" if fit_color or fit_level else "Không rõ"

    return f"""Mirror Check response:
- Micro-portrait app đề xuất mạnh nhất: {score_matched}
- Học sinh chọn micro-portrait nào: {student_choice}
- Câu học sinh tick/highlight là đúng nhất (PHẦN ĐÚNG): {highlight}
- Câu học sinh nói KHÔNG giống mình / KHÔNG phải lựa chọn của em (PHẦN BỊ LOẠI, không phải một phần trong lựa chọn ở trên): {mismatch}
- Mirror Fit sơ bộ: {fit_combined}

LƯU Ý QUAN TRỌNG: dòng "PHẦN BỊ LOẠI" ở trên là portrait mà học sinh nói KHÔNG giống mình —
đây KHÔNG phải một trong hai portrait tạo nên lựa chọn "pha giữa" của học sinh (nếu có).
Chỉ dùng đúng 2 chữ cái đã nêu trong dòng "Học sinh chọn micro-portrait nào" ở trên khi mô tả
lựa chọn của học sinh — không tự suy ra hoặc thay thế bằng chữ cái nào khác."""


def build_prompt(
    student_info: Dict[str, Any],
    scores: Dict[str, Any],
    mirror_check: Optional[Dict[str, Any]] = None,
    has_transcript_files: bool = False,
) -> str:
    sop           = _read_file(SOP_FILENAME)
    master_router = _read_file(MASTER_ROUTER_FILENAME)

    route = _detect_route(student_info.get("direction", ""))

    extra_branch_doc = ""
    if route == "B":
        extra_branch_doc = "\n\n---\n\n# TÀI LIỆU BỔ SUNG — SIÊU PROMPT 2.5 (TRONG NƯỚC)\n\n" + _read_file(TRONG_NUOC_FILENAME)

    # Two independent ways transcript info can arrive:
    #   1. has_transcript_files=True — actual file(s) (PDF/image) are attached
    #      as separate document/image content blocks in the API call itself
    #      (see _build_transcript_content_blocks) — Claude reads them
    #      directly, this is just a text pointer telling it they're there.
    #   2. student_info["transcript"] — a plain text description (legacy
    #      fallback, used if no real files were ever wired in for this call).
    # Both can't really apply at once in practice, but handle gracefully
    # either way rather than assuming only one path is ever used.
    if has_transcript_files:
        transcript_line = (
            "\nHọc bạ (bằng chứng minh họa, không tự chốt hướng): đã đính kèm bên dưới "
            "dưới dạng file/ảnh — đọc trực tiếp nội dung file để lấy thông tin điểm số, "
            "môn mạnh/môn yếu. Không bịa thêm nếu file khó đọc; nếu vậy, ghi rõ trong "
            "phần audit rằng cần xác nhận lại thủ công."
        )
    else:
        transcript = student_info.get("transcript")
        transcript_line = f"\nHọc bạ (bằng chứng minh họa, không tự chốt hướng): {transcript}" if transcript else ""

    mirror_check_block = _build_mirror_check_block(mirror_check)

    # Computed here, deterministically, rather than left for the model
    # to guess — this was the actual root cause of two separate issues
    # found in testing: the model sometimes left a literal
    # "[Ngày phát hành]" placeholder bracket unfilled, and other times
    # fabricated an inconsistent/wrong date (e.g. "2025", "Tháng 11,
    # 2024") — because it was never actually GIVEN a real date value to
    # use, only told in the instructions that this field should exist
    # on the title page. Every other title-page field comes from real
    # student_info data; this one now does too.
    release_date_str = datetime.now(ZoneInfo("Asia/Ho_Chi_Minh")).strftime("%d/%m/%Y")

    filled_fields = f"""
THÔNG TIN ĐIỀN VÀO MASTER ROUTER (thay cho các trường [Điền tên học sinh] v.v. ở trên):

Học sinh: {student_info.get('name', '')}
Lớp / tình trạng học tập: {student_info.get('grade', '')}
Bối cảnh đích đến: {student_info.get('direction', '')}
Dự định sau THPT: {student_info.get('after_school', '')}
Trường hiện tại: {student_info.get('school', '')}
Môn hợp vibe: {student_info.get('fav_subjects', '')}
Hoạt động yêu thích: {student_info.get('fav_activities', '')}
Tình trạng case: Test lần đầu
File đầu vào: File test (đã chấm điểm bên dưới){transcript_line}
Format cần viết theo: format An Du, đầy đủ Student Snapshot + Executive Summary + Consultant Note
Ngày phát hành báo cáo (dùng ĐÚNG giá trị này cho dòng "Ngày phát hành:" ở trang bìa, không tự
  bịa hoặc để trống): {release_date_str}

{mirror_check_block}

DỮ LIỆU ĐÃ CHẤM ĐIỂM (không bịa thêm, không tính lại — đã verify):

MBTI: type={scores['mbti']['type']}, gap_avg={scores['mbti']['gap_avg']}, clarity={scores['mbti']['clarity']}
  Note: {scores['mbti']['note']}
  Axes: {scores['mbti']['axes']}

Holland: groups={scores['holland']['groups']}, top3={scores['holland']['top3']}

OCEAN: {scores['ocean']['groups']}

SSS: score={scores['sss']['score']}, level={scores['sss']['interpretation']}
"""

    task = """
NHIỆM VỤ:
Dùng đúng quy trình ở tài liệu "QUY TRÌNH CHỐT CASE" và format ở "SIÊU PROMPT 3.0 MASTER
ROUTER" phía trên (cùng với tài liệu bổ sung Trong Nước nếu có) để viết báo cáo hướng nghiệp
cá nhân đầy đủ cho học sinh này. Điền đúng route, chạy đúng chuỗi suy luận, không nhảy bước.
Không bịa tên trường/chương trình/số liệu nếu không chắc. Không cắt ngắn để tiết kiệm độ dài
— đây là báo cáo gửi gia đình thật.

Trước khi viết bất kỳ kết luận ngành/major nào, PHẢI đi qua MIRROR CHECK GATE (6 câu hỏi bắt
buộc) theo đúng yêu cầu trong SIÊU PROMPT 3.0 MASTER ROUTER, dùng dữ liệu Mirror Check response
ở trên. Nếu Mirror Fit là Cam hoặc Đỏ, không được viết ngành/major theo giọng chốt cứng —
phải nêu rõ vùng cần Quest kiểm chứng trước khi khẳng định.

YÊU CẦU BẮT BUỘC VỀ CẤU TRÚC ĐẦU RA (áp dụng thêm, ngoài quy trình ở trên):

1. **KHÔNG in 4 chữ MBTI (ví dụ "ENTP") ở bất kỳ đâu trong báo cáo — kể cả trong bảng dữ
   liệu/bảng điểm tổng hợp.** Đây là tài liệu gửi thẳng cho gia đình — không có phần nội bộ
   riêng nữa (xem quy tắc 10). Chỉ tả tính cách bằng ngôn ngữ thường ("em có xu hướng...",
   "em hợp kiểu..."), không gắn nhãn 4 chữ cái ở bất kỳ đâu trong toàn bộ output.
   **QUAN TRỌNG:** quy tắc này áp dụng cho MỌI hình thức trình bày, kể cả nếu bạn tạo một
   bảng "điểm test gốc" hoặc "tổng hợp kết quả" — cột "Kết quả"/"Giá trị" cho dòng MBTI trong
   bảng đó KHÔNG được viết 4 chữ cái, chỉ được viết mô tả bằng lời (ví dụ: "Hướng nội, cảm
   nhận, tình cảm, có kế hoạch — độ rõ nghiêng vừa", không viết "ISFJ, độ rõ nghiêng vừa").
   Lỗi này đã xảy ra thật trong một lần test trước — hãy tự kiểm tra kỹ trước khi xuất bảng
   nào có nhắc đến MBTI.

2. **Xếp hạng major family/vùng nghề ưu tiên #1 phải tương ứng với mã Holland điểm cao nhất**
   trong Top 3, trừ khi có lý do rõ từ OCEAN/SSS/bối cảnh để hạ xuống — nếu đảo thứ tự, phải
   nêu lý do cụ thể ngay trong đoạn phân tích Holland (mục 4), không có phần audit riêng để
   đẩy lý do đó sang.

3. **Tên đọc riêng phải LÀ chính tiêu đề đó, không phải placeholder, VÀ phải là heading cấp 2
   (H2) riêng, đứng NGAY ĐẦU tiên trong Phần A** — trước cả mục `a) Em là kiểu người như thế
   nào`, không chôn nó làm heading cấp 3 bên trong một mục con khác. Viết heading là chính
   cụm từ đã nghĩ ra (ví dụ `## Người vẽ ý tưởng có phanh`), TUYỆT ĐỐI không viết heading kiểu
   `## [TÊN ĐỌC RIÊNG]` rồi mới ghi tên thật ở dòng dưới. Cụm từ ngắn (4-6 chữ), giàu hình
   ảnh, tóm gọn cách học sinh tạo giá trị (không phải tên nghề, không phải nhãn tính cách). Sau
   tên, viết 2-3 câu giải thích, làm rõ tên này không ép học sinh vào một nghề cố định. Đây là
   điều đầu tiên người đọc thấy khi mở Phần A — phải nổi bật, không bị chôn trong một mục con.

4. **Mỗi major family/vùng nghề đề xuất phải được đào sâu đầy đủ**, không chỉ liệt kê tên:
   - Vì sao hợp (nối tới Holland/OCEAN/SSS/bối cảnh cụ thể của học sinh này)
   - Ngành này học gì (course content thực tế, không bịa tên trường cụ thể)
   - Việc thường gặp sau khi ra trường (5-8 việc cụ thể)
   - Vai trò nghề có thể hướng tới (3-5 chức danh cụ thể)
   - Mức ưu tiên (Rất cao / Cao / Khá cao / Có điều kiện / Trung bình) kèm lý do ngắn — viết
     dòng này dạng `**Mức ưu tiên:** Rất cao — <lý do>` bằng CHỮ THUẦN, không thêm bất kỳ
     emoji/icon/ký hiệu trang trí nào (không ⭐, không 🎯, không bất kỳ ký tự Unicode
     trang trí nào) trước hoặc sau.

5. **Phải có "Application Story Themes"** trong phần đầy đủ — 3-4 trục câu chuyện cho personal
   statement, mỗi trục 2-3 câu, gắn cụ thể vào dữ liệu/bối cảnh thật của học sinh này (không
   viết chung chung kiểu "em rất thích giúp người").

6. **Phải có "Hồ sơ nên có"** trong phần đầy đủ — personal statement nên xoay quanh chủ đề gì,
   loại portfolio/hoạt động nên có, loại project nên làm, hướng thư giới thiệu nên nhấn vào
   điều gì.

7. **Phải có mục "Lời kết gửi phụ huynh"** ở cuối — heading viết thường, KHÔNG có dấu ngoặc
   vuông bao quanh (viết `## Lời kết gửi phụ huynh`, không viết `## [LỜI KẾT GỬI PHỤ HUYNH]`).
   Đoạn ngắn (4-6 câu) tóm gọn tinh thần báo cáo: đây là bản đồ mở không phải bản án; nhấn lại
   điểm mạnh cốt lõi; nhắc gia đình tránh đẩy con theo hướng ngược với dữ liệu.

8. **Phải có mục "Mirror Check / Self-confirmation"** trong phần phân tích chi tiết, nêu rõ:
   Mirror Fit, vùng lệch nếu có, ảnh hưởng đến độ tin cậy của report, và Quest theo dõi
   tương ứng. Đây là báo cáo family-facing duy nhất — không có báo cáo nội bộ riêng để tách
   thông tin này ra, nên phải viết đủ chi tiết ngay tại đây bằng giọng phù hợp để phụ huynh
   đọc được (không dùng thuật ngữ kỹ thuật khô như "Mirror Mismatch Risk" trần trụi — diễn
   giải bằng câu bình thường).

9. **TUYỆT ĐỐI KHÔNG dùng emoji hoặc icon Unicode trang trí ở BẤT KỲ ĐÂU** trong toàn bộ
   output — không trong heading, không trong bullet, không ở đầu/cuối câu. Chỉ dùng chữ và
   dấu câu thông thường. Nhấn mạnh bằng **bold** hoặc *italic*, không dùng ký hiệu trang trí.

9b. **TUYỆT ĐỐI KHÔNG dùng dấu gạch ngang dài kiểu "—" (em dash) hoặc "–" (en dash) ở BẤT KỲ
    ĐÂU** trong toàn bộ output. Thay vào đó, dùng dấu phẩy, dấu hai chấm, ngoặc đơn, hoặc viết
    lại câu để không cần dấu gạch ngang. Nếu thực sự cần một dấu nối ngắn trong một cụm từ
    (ví dụ khoảng số "11-12"), chỉ dùng dấu gạch nối thường "-" (hyphen, một ký tự ngắn, khác
    với em dash/en dash).

10. **KHÔNG tạo mục "PHẦN C: CONSULTANT NOTE" hay bất kỳ phần audit/nội bộ nào ở cuối báo
    cáo.** Toàn bộ output chỉ có 2 phần lớn, không có phần thứ ba. Mọi thông tin (Holland,
    OCEAN, Career Card, META64, O*NET, VSCO, Mirror Check, cảnh báo) đều viết trực tiếp trong
    2 phần đó, bằng giọng phù hợp để gia đình đọc trực tiếp — không có nơi nào để "giấu" nội
    dung kỹ thuật ra sau một bức tường ngăn cách nữa.

11. **Đặt tên 2 phần lớn bằng chữ cái, KHÔNG dùng từ "PHẦN".** Viết heading cấp 1 (H1) là
    `# A. <tiêu đề ngắn>` và `# B. <tiêu đề ngắn>` — ví dụ `# A. Bản đọc nhanh cho em Quyên`
    và `# B. Báo cáo đầy đủ cho phụ huynh`. KHÔNG viết "PHẦN A", "Phần A –", hay bất kỳ biến
    thể nào có chữ "Phần".

12. **Trong phần A, các mục con dùng chữ cái thường + dấu ngoặc đơn** (`a)`, `b)`, `c)`...)
    làm heading cấp 2 (H2), viết liền trước tiêu đề — ví dụ `## a) Em là kiểu người như thế
    nào`. Trong phần B, các mục lớn tiếp tục đánh số 1., 2., 3... như đã quy định ở các mục
    trên (giữ nguyên cách đánh số hiện tại của phần B, không đổi). **Ngoại lệ:** tên đọc riêng
    (mục 3) và "Tổng quan nhanh" (mục 17) đứng TRƯỚC mục `a)` đầu tiên và KHÔNG dùng chữ cái —
    chỉ các mục con phân tích chi tiết SAU đó mới bắt đầu đánh chữ cái từ `a)`.

13. **Mỗi bảng markdown phải có một dòng caption in đậm ngay PHÍA DƯỚI bảng** (không phải
    phía trên), dạng `**Bảng N: <tên bảng>**`, với N là số thứ tự bảng tăng dần xuyên suốt
    toàn bộ báo cáo (không reset theo từng phần). Thứ tự đúng: bảng trước, dòng caption ngay
    sau (cách 1 dòng trống). Ví dụ:
    ```
    | Cột 1 | Cột 2 |
    |---|---|
    | A | B |

    **Bảng 1: O*NET Role Expansion**
    ```

14. **Chèn đúng 2 marker sau, mỗi marker trên một dòng riêng, không có gì khác trên dòng đó:**
    - Marker `[TOC]` — đặt ngay sau khối thông tin tiêu đề (Học sinh/Lớp/Route/Ngày phát
      hành) và TRƯỚC heading `# A. ...`. Đây sẽ được thay bằng mục lục thật.
    - Marker `[PAGEBREAK]` — đặt ngay TRƯỚC heading `# B. ...` (để phần B bắt đầu ở trang
      mới). KHÔNG đặt marker `[PAGEBREAK]` ở bất kỳ chỗ nào khác.
    Không giải thích hay nhắc đến 2 marker này trong nội dung — chúng chỉ là điểm đánh dấu kỹ
    thuật, sẽ được thay thế tự động, không hiển thị dạng chữ trong bản cuối.

15. **Trang đầu tiên (trước marker `[TOC]`) phải trình bày trang trọng, giống trang bìa**:
    tiêu đề báo cáo là `# BÁO CÁO HƯỚNG NGHIỆP CÁ NHÂN`, theo sau là khối thông tin học sinh
    (Học sinh / Lớp-Trường / Định hướng / Route case / Ngày phát hành) viết dạng
    `**Nhãn:** Giá trị` mỗi dòng — không thêm nội dung phân tích nào khác ở trang này. Dòng
    "Ngày phát hành" PHẢI dùng đúng giá trị đã cho ở mục "Ngày phát hành báo cáo" bên dưới —
    KHÔNG để trống, KHÔNG viết placeholder dạng "[Ngày phát hành]", KHÔNG tự đoán ngày khác.

16. **Mọi con số điểm test khi nhắc đến trong bài viết phải kèm theo thang điểm tối đa**,
    dạng `X/Y`, để người đọc hình dung được vị trí trên thang — không viết số điểm trần trụi
    một mình. Ví dụ: OCEAN Openness "4.25/5" (không chỉ viết "4.25"); Holland "A (45/50)"
    (không chỉ viết "A (45)"); SSS "2.83/5". Áp dụng cho MỌI con số điểm xuất hiện trong toàn
    bộ báo cáo, kể cả trong bảng.

17. **Phải có mục "Tổng quan nhanh" (heading `## Tổng quan nhanh`), đứng NGAY SAU tên đọc
    riêng (mục 3) và TRƯỚC mục `a) Em là kiểu người như thế nào`.** Đây là bản tóm tắt cực
    ngắn, để một phụ huynh bận rộn đọc lướt 30 giây là nắm được ý chính, trước khi đọc chi
    tiết ở các phần sau. Bắt buộc có đủ 3 phần, mỗi phần chỉ 2-4 câu hoặc vài gạch đầu dòng,
    KHÔNG lặp lại toàn bộ nội dung chi tiết đã có ở Phần B (chỉ tóm tắt, chi tiết đầy đủ vẫn
    nằm ở các mục sau như đã quy định):
    - **Kết luận định hướng:** 1-2 câu nêu hướng chính đang nổi lên từ dữ liệu.
    - **Vùng ngành nên ưu tiên** (danh sách ngắn, xếp theo thứ tự ưu tiên, mỗi dòng 1 câu):
      liệt kê nhanh các vùng ngành sẽ được đào sâu ở Phần B, kèm mức ưu tiên ngắn gọn.
    - **Vùng nên tránh hoặc cân nhắc kỹ** (danh sách ngắn): liệt kê nhanh những hướng KHÔNG
      phù hợp hoặc cần thận trọng, sẽ được giải thích đầy đủ ở mục cảnh báo phía sau.
"""

    return (
        "# TÀI LIỆU 1 — QUY TRÌNH CHỐT CASE\n\n" + sop +
        "\n\n---\n\n# TÀI LIỆU 2 — SIÊU PROMPT 3.0 MASTER ROUTER\n\n" + master_router +
        extra_branch_doc +
        "\n\n---\n\n# TÀI LIỆU 3 — DỮ LIỆU HỌC SINH VÀ NHIỆM VỤ\n\n" + filled_fields +
        "\n\n" + task
    )


_REFERENCE_DOCX_CACHE: Path | None = None


def _force_font_no_theme(style, font_name: str) -> None:
    """
    Set a style's font to a literal font name AND strip any theme font
    references (asciiTheme/hAnsiTheme/eastAsiaTheme/cstheme). Heading
    styles in Word templates (including pandoc's default reference doc)
    often carry ONLY a theme reference with no literal font at all —
    setting style.font.name alone adds a literal override but leaves
    the theme reference in place, and renderers can still follow the
    theme instead. Verified by inspecting the actual generated XML:
    without this, Heading styles kept asciiTheme="majorHAnsi" etc. and
    rendered in the theme's font (Calibri-like) instead of Arial.
    """
    style.font.name = font_name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return
    for theme_attr in ["asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"]:
        attr_qn = qn(f"w:{theme_attr}")
        if attr_qn in rFonts.attrib:
            del rFonts.attrib[attr_qn]
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _get_reference_docx() -> Path:
    """
    Build (once, cached) a .docx with Arial set as the default font for
    Normal and Heading styles, used as pandoc's --reference-doc.

    Starts from pandoc's OWN default reference doc (via
    `pandoc --print-default-data-file reference.docx`), not a blank
    python-docx Document() — a blank document lacks proper numbering.xml
    definitions, which silently breaks bullet/numbered list rendering.
    """
    global _REFERENCE_DOCX_CACHE
    if _REFERENCE_DOCX_CACHE and _REFERENCE_DOCX_CACHE.exists():
        return _REFERENCE_DOCX_CACHE

    tmp_dir = Path(tempfile.gettempdir())
    pandoc_default = tmp_dir / "an_du_pandoc_default.docx"
    path = tmp_dir / "an_du_reference.docx"

    result = subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        capture_output=True
    )
    if result.returncode != 0:
        raise RuntimeError(f"Failed to get pandoc default reference doc: {result.stderr}")
    pandoc_default.write_bytes(result.stdout)

    doc = Document(str(pandoc_default))

    # Iterate by name rather than dict-style lookup (doc.styles["Heading 1"]
    # raised KeyError on this specific template despite the style being
    # present when iterated — a latent-style quirk in pandoc's reference doc).
    target_names = {"Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"}
    for style in doc.styles:
        if style.name in target_names:
            _force_font_no_theme(style, "Arial")
            # Black text everywhere, not the reference doc's theme blue —
            # this is a style-level default; _normalize_fonts_everywhere()
            # (run after pandoc conversion) additionally forces this at the
            # individual run level too, since not everything reliably
            # inherits style-level color (e.g. table cell text).
            style.font.color.rgb = RGBColor(0, 0, 0)
            if style.name == "Normal":
                style.font.size = Pt(11)
                # Body text justified, not left-aligned — headings are left
                # untouched (they inherit their own alignment, which stays
                # left by default and should — justify only makes sense for
                # multi-line body paragraphs).
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if style.name.startswith("Heading"):
                # Extra breathing room between a paragraph and the next
                # heading, so sections don't feel cluttered.
                style.paragraph_format.space_before = Pt(18)

    # Explicitly set A4 portrait page size + 1-inch margins on the
    # default section. Pandoc's own default reference doc declares
    # NEITHER a page size nor margins at all (confirmed by inspection —
    # section.page_width etc. all return None), leaving it entirely to
    # whichever renderer's own implicit default happens to apply. This
    # is the likely root cause of two separate reported issues: (1) a
    # crash in _get_page_content_width_dxa needing an explicit width to
    # compute fixed table sizing, and (2) some pages rendering as
    # landscape — different renderers (Word, LibreOffice, Google Docs)
    # can guess differently when nothing is declared, especially across
    # the multiple sections _center_title_page introduces. Vietnamese
    # audience → A4, not US Letter.
    section = doc.sections[0]
    section.page_width = Mm(210)   # A4
    section.page_height = Mm(297)  # A4
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    doc.save(path)
    _REFERENCE_DOCX_CACHE = path
    return path


_TOC_FIELD_OOXML = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
<w:sdt>
  <w:sdtPr>
    <w:docPartObj>
      <w:docPartGallery w:val="Table of Contents"/>
      <w:docPartUnique/>
    </w:docPartObj>
  </w:sdtPr>
  <w:sdtContent>
    <w:p>
      <w:pPr><w:pStyle w:val="TOCHeading"/></w:pPr>
      <w:r><w:t>Mục lục</w:t></w:r>
    </w:p>
    <w:p>
      <w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>
      <w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>(Nhấn phải chuột chọn "Update Field" hoặc phím F9 để cập nhật mục lục)</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>
</w:sdtContent>
</w:sdt>
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""

_PAGEBREAK_OOXML = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""

# Matches common decorative Unicode emoji ranges — used as a backstop in
# case the model doesn't fully comply with the "no emoji" prompt rule.
# Deliberately does NOT strip ordinary Vietnamese diacritics or CJK
# characters, which live in entirely different Unicode blocks.
_EMOJI_PATTERN = re.compile(
    "["
    "\U0001F300-\U0001FAFF"  # symbols & pictographs, emoticons, transport, supplemental symbols
    "\U00002600-\U000027BF"  # misc symbols, dingbats (includes ⭐ U+2B50 is outside this — added below)
    "\U00002B00-\U00002BFF"  # misc symbols and arrows (covers ⭐ U+2B50)
    "\U0001F1E6-\U0001F1FF"  # regional indicator symbols (flags)
    "\U0000FE0F"             # variation selector-16 (emoji presentation)
    "]+",
    flags=re.UNICODE,
)


def _strip_trailing_consultant_section(text: str) -> str:
    """
    Safety net: the prompt now instructs the model to never generate a
    separate consultant/audit section (Phần C), but LLM instruction-
    following isn't 100% guaranteed — this truncates anything from a
    recognizable "Part C" HEADING onward if one slips through anyway,
    so a prompt-compliance miss can't leak internal-only content (MBTI
    codes, raw audit notes) into the family-facing document.

    CRITICAL: every pattern below must be anchored to an actual
    markdown heading line (starting with #) — NOT free-floating text
    anywhere in the document. An earlier version matched the bare
    phrase "PHẦN C" anywhere in the text, which caught a real, severe
    false positive in testing: the Mirror Check discussion legitimately
    contains sentences like "em thấy có phần C (thích tương tác...)"
    (referring to Portrait C), which matched that pattern at character
    ~1850 of a 33,000-character report and silently truncated ~95% of
    an otherwise good generation down to a single page. Never again —
    every pattern here requires the match to be on its own heading
    line, not just present somewhere in prose.
    """
    markers = [
        r"^#{1,3}\s*C\.\s",                  # "# C. ..." or "## C. ..." — matches the letter-heading convention (rules 11/12) if the model reverts to a forbidden 3rd top-level part
        r"^#{1,3}\s*.*CONSULTANT\s+NOTE",     # a heading line containing "CONSULTANT NOTE"
        r"^#{1,3}\s*.*PHẦN\s+C\b",            # a heading line containing the legacy "PHẦN C" naming, ONLY if it's actually a heading — not prose mentioning "phần C" (e.g. "em thấy có phần C giống mình")
        r"\[AUDIT\s+NỘI\s+BỘ\]",              # distinctive bracketed marker — specific enough to be safe as a free-floating match, unlikely to ever occur by coincidence in natural prose
    ]
    earliest_cut = None
    for pattern in markers:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match and (earliest_cut is None or match.start() < earliest_cut):
            earliest_cut = match.start()
    if earliest_cut is not None:
        print(f"WARNING: Trailing consultant/audit section detected and stripped "
              f"(model didn't fully follow the 'no Part C' instruction) — "
              f"cut at character {earliest_cut}")
        return text[:earliest_cut].rstrip()
    return text


def _replace_markers(text: str) -> str:
    """Replace the [TOC] and [PAGEBREAK] markers the prompt instructs the
    model to emit with their real raw-OOXML equivalents."""
    text = text.replace("[TOC]", _TOC_FIELD_OOXML)
    text = text.replace("[PAGEBREAK]", _PAGEBREAK_OOXML)
    return text


def _ensure_pagebreak_before_part_b(text: str) -> str:
    """
    Guarantees a page break exists immediately before the "# B. ..."
    heading, regardless of whether the model included the [PAGEBREAK]
    marker rule 14 instructs it to place there.

    Found in testing (token HN-2026-0011): the model sometimes omits
    the marker entirely — not malformed or misplaced, genuinely absent
    from the output, confirmed by searching for literal "PAGEBREAK"
    text and finding none. Rather than relying on stronger prompt
    wording (which only improves the odds, never guarantees it — the
    same lesson from every other mechanical-formatting fix this
    session), this finds the "# B." heading directly in the raw
    markdown and inserts the page break there unconditionally. Runs
    AFTER _replace_markers(), so any [PAGEBREAK] marker the model DID
    include has already become a real OOXML block by this point — this
    only needs to check for the heading itself, not worry about
    double-inserting if the marker was also present, since the marker
    (when present) is typically placed on its own line just before the
    heading, not immediately adjacent to it, and this looks for the
    heading pattern specifically, not the marker.
    """
    match = re.search(r"^#\s*B\.\s", text, re.MULTILINE)
    if not match:
        print("WARNING: No '# B.' heading found — cannot guarantee a page break "
              "before Part B. This likely means Part B is missing entirely, a "
              "more serious issue than a missing page break.")
        return text

    insert_pos = match.start()
    # Avoid inserting a second, redundant page break immediately above
    # this heading if one is already there from the model's own
    # [PAGEBREAK] marker (now expanded to _PAGEBREAK_OOXML text) sitting
    # directly adjacent, just before it.
    preceding_text = text[:insert_pos]
    if preceding_text.rstrip().endswith(_PAGEBREAK_OOXML.rstrip()):
        return text

    return preceding_text + "\n" + _PAGEBREAK_OOXML + "\n\n" + text[insert_pos:]


def _strip_stray_emoji(text: str) -> str:
    """Backstop for the 'no emoji' prompt rule — removes any that slip
    through despite the instruction, rather than relying on compliance
    alone."""
    return _EMOJI_PATTERN.sub("", text)


def _fix_over_escaped_bold_markers(text: str) -> str:
    """
    Fixes a real content quirk found in testing (token HN-2026-0011,
    2026-07-14): the model correctly escapes the literal asterisk in
    "O*NET" as "O\\*NET" (a real product name — smart to escape, since
    an unescaped "*" there risks pandoc parsing it as emphasis
    syntax), but this escaping habit occasionally bled into the
    model's OWN bold markdown syntax nearby, turning intended
    **Bảng 2: ...** captions into literal "\\*\\*Bảng 2: ...\\*\\*"
    text (visible backslash-asterisks instead of real bold formatting).

    A legitimate single-asterisk escape (O\\*NET) uses ONE backslash-
    asterisk pair. An over-escaped bold marker uses TWO consecutive
    backslash-asterisk pairs (\\*\\*), since the model tried to escape
    both asterisks of its own ** markdown syntax. This targets exactly
    that four-character sequence, leaving legitimate single escapes
    (like O\\*NET) completely untouched.
    """
    return text.replace("\\*\\*", "**")


def _strip_em_en_dashes(text: str) -> str:
    """
    Replaces every em dash (—, U+2014) and en dash (–, U+2013) with a
    plain hyphen-with-spaces (" - "). Backstop for build_prompt()'s
    rule 9b — added on request, since the model has used both dash
    variants pervasively in generated text (e.g. "Route B – Học tiếp
    trong nước") despite this being exactly the kind of stylistic
    instruction that isn't 100% reliable from prompt compliance alone,
    same reasoning as the emoji-stripping and bold-marker backstops
    above. Ordinary hyphens ("-", U+002D) are untouched — only the
    two longer dash characters are targeted.
    """
    text = text.replace("—", " - ")
    text = text.replace("–", " - ")
    # Collapse any doubled-up spacing the replacement might introduce
    # (e.g. "word — word" -> "word  -  word" -> "word - word").
    text = re.sub(r"\s+-\s+", " - ", text)
    return text


def _strip_heading_numbering(doc: Document) -> None:
    """
    Remove any list-numbering association from heading/title paragraphs.
    Best-effort fix for a small icon appearing next to headings in some
    viewers — confirmed NOT coming from style-level numbering (checked
    directly), so this may be a Google Docs-specific rendering of the
    bookmark anchors our TOC field's \\h switch creates at each heading
    (needed for the TOC's clickable links) rather than something in the
    document itself. Kept as a defensive no-cost safety net regardless.
    """
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") or p.style.name == "Title":
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    pPr.remove(numPr)


def _center_title_page(doc: Document) -> None:
    """
    Horizontally AND vertically centers the title page (everything
    before the first page break — which is the leading break emitted
    by the [TOC] marker's replacement, right after the title/student
    info block).

    Vertical centering on a single page (not the whole document)
    requires giving that content its own Word SECTION with
    w:vAlign=center — plain page breaks can't do this, they stay
    within one section. This finds the paragraph carrying the first
    manual page-break run, strips that literal break (a section break
    inserted here already forces a new page on its own), and attaches
    a cloned copy of the document's default section properties (so
    the title page keeps the same page size/margins as the rest of
    the doc) with vAlign=center added.
    """
    body = doc.element.body
    body_sectPr = body.find(qn("w:sectPr"))
    if body_sectPr is None:
        return  # unexpected reference-doc shape — bail safely, no centering rather than a crash

    boundary_found = False
    for p in doc.paragraphs:
        br = p._p.find(".//" + qn("w:br"))
        if br is not None and br.get(qn("w:type")) == "page":
            for run_el in list(p._p.findall(qn("w:r"))):
                p._p.remove(run_el)
            new_sectPr = copy.deepcopy(body_sectPr)
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), "center")
            new_sectPr.append(vAlign)
            pPr = p._p.get_or_add_pPr()
            pPr.append(new_sectPr)
            boundary_found = True
            break

    if not boundary_found:
        print("WARNING: _center_title_page could not find a page-break boundary "
              "paragraph — title page centering skipped. Report still generated normally.")
        return

    # Horizontally center every paragraph before that same boundary.
    for p in doc.paragraphs:
        if p._p.find(".//" + qn("w:sectPr")) is not None:
            break
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _get_page_content_width_dxa(doc: Document) -> int:
    """
    Compute the usable content width (page width minus left/right
    margins) of the document's default section, in twentieths of a
    point (dxa — the unit OOXML table widths use). Used to force
    tables to a FIXED total width matching the page, rather than
    letting them auto-expand based on cell content — which is the
    likely cause of some pages rendering as landscape (a renderer's
    "table doesn't fit, try rotating the page" heuristic kicking in
    for wide tables with long cell text).

    _get_reference_docx() now explicitly sets A4 + 1in margins, so this
    should always resolve cleanly — the fallback below (same A4 + 1in
    values, computed in dxa directly) only matters if some future
    change to that function accidentally drops the explicit page size
    again, same failure mode confirmed during testing (pandoc's own
    default reference doc has NO page size declared at all).
    """
    section = doc.sections[0]
    width = section.page_width
    left = section.left_margin
    right = section.right_margin
    if width is None or left is None or right is None:
        print("WARNING: Section page_width/margins are None — falling back to "
              "A4 + 1in margin defaults for table width calculation.")
        return 11906 - 1440 - 1440  # A4 width dxa - 1in - 1in margins
    # .twips converts python-docx's internal EMU representation to dxa
    # (twentieths of a point) — the unit OOXML's w:tblW actually expects.
    # Must convert EACH Length to .twips individually before subtracting
    # — subtracting two Length objects returns a plain int (loses the
    # .twips conversion method), caught during testing as an
    # AttributeError on the very next line after fixing a related bug.
    return int(width.twips - left.twips - right.twips)


def _set_cell_borders(cell) -> None:
    """
    Explicitly set borders on a single table cell (all 4 edges).

    Cell-level borders take precedence over table-level borders in
    OOXML's border resolution order — table-level (tblBorders) is the
    lowest priority, cell-level (tcBorders) wins over it. Relying on
    tblBorders alone (as an earlier version of _style_tables did) can
    have some cells' borders silently overridden if pandoc's own table
    conversion already assigned conflicting cell-level border
    properties — confirmed in testing: the border between the header
    row and first data row was invisible specifically in some columns
    but not others, exactly the pattern you'd expect from a per-cell
    override winning over a table-level default in some cells but not
    all. Setting borders explicitly on every cell removes any
    ambiguity — there's nothing left for another border definition to
    win against.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def _style_tables(doc: Document) -> None:
    """
    Post-process every table: add visible grid borders on both the
    table AND every individual cell (see _set_cell_borders — cell-
    level is what actually guarantees visibility, table-level alone
    isn't reliable), bold the header row, and center the table
    horizontally on the page. Done via python-docx after conversion
    rather than fighting pandoc's reference-doc table-style-name
    resolution, which is fragile and version-dependent.

    Also forces each table to a FIXED total width matching the page's
    content width (tblLayout=fixed + explicit tblW), instead of the
    default auto-fit-to-content behavior — a wide table with long cell
    text can otherwise exceed the page width, which is the likely
    cause of some pages rendering landscape (some renderers/converters
    auto-rotate a page to fit an overflowing table rather than
    reporting an error).
    """
    content_width_dxa = _get_page_content_width_dxa(doc)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        tbl = table._tbl
        tblPr = tbl.tblPr

        # Table-level borders — kept as a baseline default, but
        # _set_cell_borders below (applied to every cell) is what
        # actually guarantees visibility everywhere, since cell-level
        # wins any conflict with this.
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tblPr.append(borders)

        # Fixed layout + explicit total width — see docstring above.
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(content_width_dxa))

        # Distribute that fixed width evenly across columns explicitly —
        # tblLayout=fixed without per-column widths can still leave a
        # renderer to guess column proportions from content.
        num_cols = len(table.columns)
        if num_cols > 0:
            col_width = content_width_dxa // num_cols
            grid = tbl.find(qn("w:tblGrid"))
            if grid is not None:
                for gridCol in grid.findall(qn("w:gridCol")):
                    gridCol.set(qn("w:w"), str(col_width))
            for row in table.rows:
                for cell in row.cells:
                    cell.width = col_width

        # Explicit per-cell borders — the actual fix (see
        # _set_cell_borders docstring for why table-level alone isn't
        # sufficient).
        for row in table.rows:
            for cell in row.cells:
                _set_cell_borders(cell)

        if table.rows:
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _style_table_captions(doc: Document) -> None:
    """
    Find caption paragraphs matching "Bảng N: ..." and explicitly
    center + bold them. These are plain body paragraphs from the
    model's markdown output (not a distinct Word style), so they need
    the same per-paragraph treatment as everything else post-processed
    here rather than relying on a style-level default.

    Handles TWO distinct failure modes found in testing, both
    resulting in a caption that doesn't render as real bold text:
      1. Normal case: pandoc correctly parsed **Bảng N: ...** as bold
         markdown — just needs centering added (bold is already real
         formatting here).
      2. Literal-asterisk case (token HN-2026-0011, later test): the
         markdown ** syntax was NOT parsed as bold at all — the
         paragraph's actual text is the literal string
         "**Bảng 2: O*NET Role Expansion**", asterisks visibly
         present. Most likely cause: the model didn't leave a blank
         line between the table and this caption, which markdown
         requires to correctly terminate a table before parsing the
         next paragraph — pandoc's parser can then fail to recognize
         the following ** as inline emphasis syntax and just treats
         it as literal text. This case needs the asterisks physically
         stripped from the run text, THEN bold+center applied,
         otherwise the caption ships with visible "**" characters.
    """
    normal_pattern = re.compile(r"^Bảng\s+\d+\s*:", re.IGNORECASE)
    literal_asterisk_pattern = re.compile(r"^\*\*\s*(Bảng\s+\d+\s*:.*?)\*\*\s*$", re.IGNORECASE)

    for p in doc.paragraphs:
        text = p.text.strip()

        if normal_pattern.match(text):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
            continue

        literal_match = literal_asterisk_pattern.match(text)
        if literal_match:
            # Strip the literal ** characters from the actual run text,
            # then apply real bold + center formatting — otherwise this
            # ships with visible asterisks instead of proper bold text.
            clean_text = literal_match.group(1).strip()
            # Clear existing runs and replace with a single clean run,
            # simplest reliable way to change run-level text content
            # via python-docx (Run objects don't expose a text setter
            # that handles removing/replacing cleanly otherwise).
            for run in list(p.runs):
                run.text = ""
            if p.runs:
                p.runs[0].text = clean_text
                p.runs[0].bold = True
            else:
                new_run = p.add_run(clean_text)
                new_run.bold = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"WARNING: Fixed a table caption with unparsed literal markdown "
                  f"asterisks: {clean_text[:50]!r}")


def _force_portrait_orientation(doc: Document) -> None:
    """
    Explicitly force every section's page orientation to portrait and
    ensure page width < height accordingly. Defensive — in case any
    section (including the title-page section created by
    _center_title_page) ends up with a landscape-flagged pgSz for any
    reason, this guarantees the final output is portrait throughout.
    """
    for section in doc.sections:
        sectPr = section._sectPr
        pgSz = sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            pgSz = OxmlElement("w:pgSz")
            sectPr.append(pgSz)
        # Remove any landscape orientation flag
        if qn("w:orient") in pgSz.attrib:
            del pgSz.attrib[qn("w:orient")]
        width = int(pgSz.get(qn("w:w")) or 0)
        height = int(pgSz.get(qn("w:h")) or 0)
        if width and height and width > height:
            # Swap so width < height (portrait) — this section's
            # dimensions were landscape; correct them.
            pgSz.set(qn("w:w"), str(height))
            pgSz.set(qn("w:h"), str(width))



def _normalize_run_element_font(r_element) -> None:
    """
    Low-level version of _normalize_run_font that operates directly on
    a raw <w:r> XML element rather than a python-docx Run wrapper —
    needed because the recursive body-wide walk in
    _normalize_fonts_everywhere() below finds runs python-docx's
    higher-level Run objects aren't conveniently constructed for
    (e.g. runs inside a w:sdt wrapper).
    """
    rPr = r_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_element.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        attr_qn = qn(f"w:{theme_attr}")
        if attr_qn in rFonts.attrib:
            del rFonts.attrib[attr_qn]
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")

    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), "000000")


def _normalize_fonts_everywhere(doc: Document) -> None:
    """
    Force Arial + black on EVERY run in the document — including ones
    doc.paragraphs/doc.tables don't reach at all.

    Confirmed by testing (token HN-2026-0011, 2026-07-14): the earlier
    version of this function only walked doc.paragraphs + doc.tables,
    which silently MISSED the TOC field's own "Mục lục" heading text —
    that content lives inside a <w:sdt> structured-document-tag wrapper
    (part of the TOC field's own required structure — see
    _TOC_FIELD_OOXML), and python-docx's doc.paragraphs property only
    enumerates <w:p> elements that are DIRECT children of the body,
    not ones nested inside a wrapper element like w:sdt. The result:
    that one heading kept the TOC style's own default Calibri+blue
    instead of getting normalized to Arial+black like everything else.

    Fixed by walking the raw XML body recursively for every <w:r>
    element regardless of nesting depth — catches paragraphs, table
    cells, AND anything inside w:sdt or other wrapper elements in one
    pass, rather than trying to enumerate every possible container
    type doc.paragraphs/doc.tables might miss.
    """
    body = doc.element.body
    for r_element in body.iter(qn("w:r")):
        _normalize_run_element_font(r_element)


def _post_process_docx(docx_path: Path) -> None:
    """
    Single load/save pass applying every python-docx-level fix:
    strip stray heading numbering, center the title page (its own
    section with vAlign=center), force portrait orientation on every
    section, style tables (borders/bold/centered/fixed-width), style
    table captions (centered/bold), and force Arial+black everywhere.
    Order matters only in that title page centering must run before
    the general paragraph loop in _normalize_fonts_everywhere touches
    those same paragraphs — doesn't conflict here since alignment and
    font/color are independent properties, but keeping this as one
    clearly-ordered pass avoids any future edit accidentally
    introducing a conflict.
    """
    doc = Document(str(docx_path))

    _strip_heading_numbering(doc)
    _center_title_page(doc)
    _force_portrait_orientation(doc)
    _style_tables(doc)
    _style_table_captions(doc)
    _normalize_fonts_everywhere(doc)

    doc.save(str(docx_path))


def _enable_toc_auto_update(docx_path: Path) -> None:
    """
    Injects <w:updateFields w:val="true"/> into word/settings.xml,
    telling Word to automatically recalculate fields (including the
    TOC) when the document is opened — instead of requiring a manual
    right-click > Update Field / F9 the first time. Done via direct
    zip manipulation (docx is a zip archive) rather than python-docx,
    which has no high-level API for this specific setting.

    Must run AFTER all python-docx-based edits (_post_process_docx)
    are already saved — python-docx's own .save() fully rewrites the
    zip archive, which would silently discard this change if applied
    before that point instead of after.
    """
    with zipfile.ZipFile(docx_path, "r") as zin:
        settings_xml = zin.read("word/settings.xml").decode("utf-8")
        other_items = [(item, zin.read(item.filename)) for item in zin.infolist()
                        if item.filename != "word/settings.xml"]

    if "<w:updateFields" not in settings_xml:
        settings_xml = re.sub(
            r"(<w:settings[^>]*>)",
            lambda m: m.group(1) + '<w:updateFields w:val="true"/>',
            settings_xml,
            count=1,
        )

    tmp_path = str(docx_path) + ".tmp"
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item, data in other_items:
            zout.writestr(item, data)
        zout.writestr("word/settings.xml", settings_xml)
    os.replace(tmp_path, docx_path)


def markdown_to_docx_base64(markdown_text: str) -> str:
    """
    Convert markdown report text to a .docx file via pandoc,
    return as base64 string ready to send over JSON.

    Pipeline:
      1. Strip any trailing consultant/audit section that slipped
         through despite the prompt instruction not to generate one.
      2. Replace [TOC] / [PAGEBREAK] markers with real raw-OOXML blocks
         (Word TOC field, page break) — see build_prompt()'s rule 14.
      3. Strip any stray emoji that slipped through despite the
         prompt's "no emoji" rule.
      4. Run pandoc (with raw_attribute enabled, required for the raw
         OOXML blocks above to pass through instead of being escaped).
      5. Post-process the resulting docx (_post_process_docx):
         - strip any stray heading numbering
         - center the title page (its own section, vAlign=center)
         - style every table (borders, bold header, centered)
         - force Arial + black on every run, including table cells
      6. Enable TOC auto-update on open (_enable_toc_auto_update) —
         zip-level settings.xml edit, must run after step 5's saves.

    Apps Script usage (decode + save to Drive — runs as real user,
    so no service-account storage quota issue):

        var bytes = Utilities.base64Decode(result.docx_base64);
        var blob  = Utilities.newBlob(bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'BC_' + token + '.docx');
        var file  = folder.createFile(blob);
    """
    markdown_text = _strip_trailing_consultant_section(markdown_text)
    markdown_text = _replace_markers(markdown_text)
    markdown_text = _ensure_pagebreak_before_part_b(markdown_text)
    markdown_text = _strip_stray_emoji(markdown_text)
    markdown_text = _fix_over_escaped_bold_markers(markdown_text)
    markdown_text = _strip_em_en_dashes(markdown_text)

    with tempfile.TemporaryDirectory() as tmp:
        md_path   = Path(tmp) / "report.md"
        docx_path = Path(tmp) / "report.docx"
        md_path.write_text(markdown_text, encoding="utf-8")

        reference_doc = _get_reference_docx()

        result = subprocess.run(
            ["pandoc", "-f", "markdown+hard_line_breaks+raw_attribute", str(md_path),
             "-o", str(docx_path), f"--reference-doc={reference_doc}"],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            raise RuntimeError(f"pandoc conversion failed: {result.stderr}")

        _post_process_docx(docx_path)
        _enable_toc_auto_update(docx_path)

        docx_bytes = docx_path.read_bytes()
        return base64.b64encode(docx_bytes).decode("utf-8")



def generate_report_stream(
    student_info: Dict[str, Any],
    scores: Dict[str, Any],
    mirror_check: Optional[Dict[str, Any]] = None,
) -> Generator[str, None, None]:
    """
    True streaming version — yields each text chunk AS IT ARRIVES from
    Anthropic, so the HTTP response itself carries continuous bytes the
    whole time generation is running. This is what actually prevents an
    idle-connection gateway timeout (Railway or otherwise): the previous
    non-streaming-HTTP version called Anthropic with streaming internally
    but still sent ZERO bytes to the client until everything finished —
    which does nothing to keep an idle-timeout-based proxy from cutting
    the connection.

    Yields the report text progressively, then a final metadata block:

        <report text, streamed chunk by chunk>

        ===METADATA===
        {"model": ..., "input_tokens": ..., "output_tokens": ..., "estimated_cost_usd": ...}

    Does NOT include docx conversion — call /webhook/markdown-to-docx
    separately with the collected text if you need the .docx file.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY environment variable not set")

    client = anthropic.Anthropic(api_key=api_key, timeout=900.0)
    prompt = build_prompt(student_info, scores, mirror_check)

    print(f"=== Starting streaming generation for {student_info.get('name', '')} "
          f"(prompt length: {len(prompt)} chars) ===")

    full_text_parts = []
    with client.messages.stream(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        for chunk in stream.text_stream:
            full_text_parts.append(chunk)
            yield chunk  # real bytes flow to the HTTP client immediately
        final_message = stream.get_final_message()

    text = "".join(full_text_parts)
    input_tokens  = final_message.usage.input_tokens
    output_tokens = final_message.usage.output_tokens
    cost = (input_tokens / 1_000_000 * 5) + (output_tokens / 1_000_000 * 25)

    print(f"=== REPORT GENERATED — name={student_info.get('name', '')} "
          f"input_tokens={input_tokens} output_tokens={output_tokens} "
          f"cost_usd={round(cost, 4)} ===")
    print("=== REPORT TEXT START ===")
    print(text)
    print("=== REPORT TEXT END ===")

    metadata = {
        "model": MODEL,
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "estimated_cost_usd": round(cost, 4),
    }
    yield "\n\n===METADATA===\n" + json.dumps(metadata, ensure_ascii=False)


def generate_report(
    student_info: Dict[str, Any],
    scores: Dict[str, Any],
    mirror_check: Optional[Dict[str, Any]] = None,
    transcript_files: Optional[list] = None,
) -> Dict[str, Any]:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY environment variable not set")

    # Explicit generous timeout (default client timeout can be too short
    # for long generations like this — Anthropic recommends streaming +
    # an explicit timeout for any request expected to run several minutes).
    client = anthropic.Anthropic(api_key=api_key, timeout=900.0)

    transcript_blocks = _build_transcript_content_blocks(transcript_files)
    prompt = build_prompt(student_info, scores, mirror_check, has_transcript_files=bool(transcript_blocks))

    # Documents/images go BEFORE the text prompt in the content list —
    # this is Anthropic's documented ordering recommendation, so Claude
    # has the actual transcript content in view before reading the
    # instructions that reference it. Falls back to the old plain-string
    # form when there are no files, functionally identical to before.
    if transcript_blocks:
        message_content = transcript_blocks + [{"type": "text", "text": prompt}]
    else:
        message_content = prompt

    print(f"=== Starting generation for {student_info.get('name', '')} "
          f"(prompt length: {len(prompt)} chars, {len(transcript_blocks)} transcript file(s)) ===")

    # Stream the response instead of a single synchronous create() call —
    # this is Anthropic's documented recommendation for long-running
    # generations (full SOP + Master Router can push this well past a
    # minute), and avoids edge cases where a long non-streaming call
    # gets dropped before the final response is assembled.
    with client.messages.stream(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": message_content}],
    ) as stream:
        for _ in stream.text_stream:
            pass  # could log incremental progress here if needed
        final_message = stream.get_final_message()

    text = "".join(block.text for block in final_message.content if block.type == "text")

    input_tokens  = final_message.usage.input_tokens
    output_tokens = final_message.usage.output_tokens
    cost = (input_tokens / 1_000_000 * 5) + (output_tokens / 1_000_000 * 25)

    # Defensive logging — print the full result to stdout (visible in
    # Railway's log viewer) the moment generation succeeds, BEFORE any
    # downstream step (docx conversion, HTTP response transport) that
    # could fail or time out. Without this, a gateway timeout after a
    # successful (and billed) Anthropic call would lose the output
    # entirely with no way to recover it.
    print(f"=== REPORT GENERATED — token={student_info.get('name', '')} "
          f"input_tokens={input_tokens} output_tokens={output_tokens} "
          f"cost_usd={round(cost, 4)} ===")
    print("=== REPORT TEXT START ===")
    print(text)
    print("=== REPORT TEXT END ===")

    try:
        docx_base64 = markdown_to_docx_base64(text)
    except Exception as e:
        docx_base64 = None  # don't fail the whole request if pandoc has an issue
        print(f"WARNING: docx conversion failed: {e}")

    return {
        "report_text": text,
        "docx_base64": docx_base64,
        "model": MODEL,
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "estimated_cost_usd": round(cost, 4),
    }


def _post_callback(callback_url: str, callback_secret: str, payload: Dict[str, Any]) -> None:
    """
    POST the finished (or failed) report result back to Apps Script's
    Web App callback URL. Uses stdlib urllib rather than requests/httpx
    to avoid adding a new dependency for a single outbound call.

    Failures here are logged, not raised — by the time this runs, the
    original HTTP request that kicked off generation is long since
    finished and has nothing to return an error to. If the callback
    itself fails (Apps Script down, wrong URL, etc.), the report is
    still fully generated and logged to stdout by generate_report()
    above — recoverable by hand from Railway logs if needed, just not
    automatically delivered to the doc.
    """
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        callback_url,
        data=data,
        headers={
            "Content-Type": "application/json; charset=utf-8",
            "X-Callback-Secret": callback_secret,
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            print(f"=== Callback POST to {callback_url} returned HTTP {resp.status} ===")
    except urllib.error.HTTPError as e:
        print(f"ERROR: Callback POST to {callback_url} failed with HTTP {e.code}: {e.read()[:500]}")
    except Exception as e:
        print(f"ERROR: Callback POST to {callback_url} failed: {e}")


def _generate_report_and_callback(
    student_info: Dict[str, Any],
    scores: Dict[str, Any],
    mirror_check: Optional[Dict[str, Any]],
    transcript_files: Optional[list],
    token: str,
    callback_url: str,
    callback_secret: str,
) -> None:
    """
    Runs generate_report() to completion, then POSTs the result (or an
    error) back to callback_url. Meant to be run in a background thread
    — see generate_report_async() below — so the original HTTP request
    that triggered this can return immediately without waiting for the
    multi-minute generation to finish.
    """
    try:
        result = generate_report(student_info, scores, mirror_check, transcript_files)
        payload = {
            "token": token,
            "status": "done",
            "report_text": result["report_text"],
            "docx_base64": result["docx_base64"],
            "model": result["model"],
            "input_tokens": result["input_tokens"],
            "output_tokens": result["output_tokens"],
            "estimated_cost_usd": result["estimated_cost_usd"],
            # Echoed inside the body (not just the X-Callback-Secret header)
            # because Apps Script's doPost() cannot reliably read custom
            # request headers across all runtime versions — verifying
            # against a body field is more robust than relying on headers
            # actually surviving the trip.
            "callback_secret_echo": callback_secret,
        }
    except Exception as e:
        print(f"ERROR: Report generation failed for token {token}: {e}")
        payload = {
            "token": token,
            "status": "error",
            "error": str(e),
            "callback_secret_echo": callback_secret,
        }

    _post_callback(callback_url, callback_secret, payload)


def generate_report_async(
    student_info: Dict[str, Any],
    scores: Dict[str, Any],
    mirror_check: Optional[Dict[str, Any]],
    token: str,
    callback_url: str,
    callback_secret: str,
    transcript_files: Optional[list] = None,
) -> None:
    """
    Starts report generation in a background thread and returns
    immediately — does NOT wait for generation to finish. The caller
    (the /webhook/generate-report-async endpoint) should return a
    "started" response to its own caller right after calling this.

    When generation finishes (successfully or not), the result is
    POSTed to callback_url — see _generate_report_and_callback().
    """
    thread = threading.Thread(
        target=_generate_report_and_callback,
        args=(student_info, scores, mirror_check, transcript_files, token, callback_url, callback_secret),
        daemon=True,
    )
    thread.start()

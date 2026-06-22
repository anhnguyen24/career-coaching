"""
server/services/docs.py — Generate consultant report as a Google Doc

Approach:
1. Build a .docx file in memory using python-docx
2. Upload it to Google Drive with mimeType set to Google Docs format
   → Drive automatically converts it to a native Google Doc on upload
3. Move it into Huong Nghiep/Submission folder
4. Return the doc URL

This avoids the verbose Google Docs API request format entirely —
only Drive API scope is needed.

TODO: verify that survey_v2.json question wording matches the live
Google Form exactly. The doc now renders question text from the JSON,
not from the form, so any future wording edits must be made in both
places (or the JSON should become the single source of truth and the
form description should be regenerated from it via form_deployer.py).

Required environment variable:
    GOOGLE_SERVICE_ACCOUNT_JSON — full contents of the service account
    credentials JSON file (paste as-is into Railway env var)

Required Drive sharing:
    The "Huong Nghiep" folder (and its "Submission" subfolder) must be
    shared with the service account email as Editor.
"""

import io
import json
import os
from datetime import datetime
from typing import Any, Dict, List, Tuple

from dateutil import parser as date_parser
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

SCOPES = ["https://www.googleapis.com/auth/drive"]

DRIVE_FOLDER_NAME = "Huong Nghiep"
SUBFOLDER_NAME    = "Submission"

# Colors
BLUE       = RGBColor(0x1A, 0x73, 0xE8)
GRAY       = RGBColor(0x55, 0x55, 0x55)
ORANGE     = RGBColor(0xE3, 0x74, 0x00)
LIGHT_GRAY = "F8F9FA"

# Student info columns in response_row, 0-based (matches Form Responses 2 layout)
STUDENT_INFO_FIELDS = [
    "timestamp", "name", "token", "dob", "gender", "grade",
    "school_year", "school", "city", "email", "phone",
    "direction", "after_school", "fav_subjects", "fav_activities", "commitment",
]


def extract_student_info(row: List[Any]) -> Dict[str, Any]:
    """Build a student_info dict from the raw response row."""
    return {
        field: (row[i] if i < len(row) else "")
        for i, field in enumerate(STUDENT_INFO_FIELDS)
    }


class DocGenerator:

    def __init__(self, survey: dict):
        self._survey = survey
        self._questions_by_number = {
            q["number"]: q
            for test in survey["tests"]
            for q in test["questions"]
        }
        self._scoring_groups = {
            ts["test_id"]: {g["id"]: g["name"] for g in ts["groups"]}
            for ts in survey["scoring"]["tests"]
        }
        self._drive = self._build_drive_service()

    # ----------------------------------------------------------
    # Public
    # ----------------------------------------------------------

    def generate(
        self,
        token: str,
        student_info: Dict[str, Any],
        answers: Dict[int, int],
        scores: Dict[str, Any],
    ) -> str:
        """Build the doc, upload it, return the doc URL."""
        doc = self._build_document(token, student_info, answers, scores)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        folder_id = self._get_submission_folder_id()
        date_str  = self._format_date(student_info.get("timestamp", ""))
        name      = student_info.get("name", "Không có tên")
        filename  = f"BC_{token}_{name}_{date_str}"

        file_id = self._upload_as_google_doc(buf, filename, folder_id)
        return f"https://docs.google.com/document/d/{file_id}/edit"

    # ----------------------------------------------------------
    # Document building
    # ----------------------------------------------------------

    def _build_document(
        self,
        token: str,
        student_info: Dict[str, Any],
        answers: Dict[int, int],
        scores: Dict[str, Any],
    ) -> Document:
        doc = Document()
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10.5)

        self._add_title(doc)
        self._add_student_info(doc, token, student_info)
        self._add_horizontal_rule(doc)

        for test in self._survey["tests"]:
            self._add_test_section(doc, test, answers)
            self._add_horizontal_rule(doc)

        self._add_results_section(doc, scores)
        self._add_horizontal_rule(doc)
        self._add_footer(doc, token)

        return doc

    def _add_title(self, doc: Document):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BÁO CÁO KHẢO SÁT HƯỚNG NGHIỆP GENZ")
        run.bold = True
        run.font.size = Pt(20)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Dành cho tư vấn viên — Nội bộ")
        run2.italic = True

    def _add_student_info(self, doc: Document, token: str, info: Dict[str, Any]):
        self._add_heading1(doc, "THÔNG TIN HỌC SINH")

        self._add_heading2(doc, "A. Thông tin cơ bản")
        self._add_field(doc, "Token", token)
        self._add_field(doc, "Họ và tên", info.get("name", ""))
        self._add_field(doc, "Ngày sinh", str(info.get("dob", "")))
        self._add_field(doc, "Giới tính", info.get("gender", ""))
        self._add_field(doc, "Lớp", info.get("grade", ""))
        self._add_field(doc, "Năm học", info.get("school_year", ""))
        self._add_field(doc, "Trường", info.get("school", ""))
        self._add_field(doc, "Thành phố/Tỉnh", info.get("city", ""))

        self._add_heading2(doc, "B. Thông tin liên hệ")
        self._add_field(doc, "Email", info.get("email", ""))
        self._add_field(doc, "SĐT/Zalo", info.get("phone", ""))

        self._add_heading2(doc, "C. Bối cảnh định hướng")
        self._add_field(doc, "Quan tâm hướng", info.get("direction", ""))
        self._add_field(doc, "Dự định sau THPT", info.get("after_school", ""))
        self._add_field(doc, "Môn hợp vibe", info.get("fav_subjects", ""))
        self._add_field(doc, "Hoạt động yêu thích", info.get("fav_activities", ""))
        self._add_field(doc, "Cam kết", info.get("commitment", ""))

    def _add_test_section(self, doc: Document, test: dict, answers: Dict[int, int]):
        self._add_heading1(doc, test["name"])

        group_names = self._scoring_groups.get(test["id"], {})

        for axis, questions in self._group_consecutive_by_axis(test["questions"]):
            label = group_names.get(axis, axis)
            qrange = f"Câu {questions[0]['number']}–{questions[-1]['number']}"
            self._add_heading2(doc, f"{axis} — {label} ({qrange})")
            for q in questions:
                answer = answers.get(q["number"], "—")
                self._add_rating_answer(doc, q["number"], q["text"], answer)

    def _add_results_section(self, doc: Document, scores: Dict[str, Any]):
        self._add_heading1(doc, "KẾT QUẢ ĐIỂM SỐ")

        mbti    = scores["mbti"]
        holland = scores["holland"]
        ocean   = scores["ocean"]
        sss     = scores["sss"]

        # MBTI table
        self._add_heading2(doc, "TEST 1 — MBTI")
        rows = [["Trục", "Điểm TB nhóm A", "Điểm TB nhóm B", "Gap", "Kết quả"]]
        axis_labels = {
            "EI": ("E (Hướng ngoại)", "I (Hướng nội)"),
            "SN": ("S (Thực tế)", "N (Trực giác)"),
            "TF": ("T (Lý trí)", "F (Cảm xúc)"),
            "JP": ("J (Nguyên tắc)", "P (Linh hoạt)"),
        }
        for axis_key, axis in mbti["axes"].items():
            ids = list(axis["scores"].keys())
            a_id, b_id = ids[0], ids[1]
            label_a, label_b = axis_labels.get(axis_key, (a_id, b_id))
            winner_label = label_a if axis["winner"] == a_id else label_b
            rows.append([
                f"{a_id} / {b_id}",
                str(axis["scores"][a_id]),
                str(axis["scores"][b_id]),
                str(axis["gap"]),
                winner_label,
            ])
        self._add_table(doc, rows)

        self._add_bold_paragraph(doc, f"→ Kiểu tính cách MBTI: {mbti['type']}", size=13)
        self._add_italic_paragraph(doc, f"→ Độ rõ: {mbti['clarity']} (Gap TB: {mbti['gap_avg']})", color=GRAY)
        self._add_italic_paragraph(doc, f"→ Lưu ý: {mbti['note']}", color=ORANGE)

        # Holland table
        self._add_heading2(doc, "TEST 2 — HOLLAND")
        holland_names = self._scoring_groups.get("holland", {})
        rows = [["Nhóm", "Điểm", "Tối đa"]]
        for gid, score in holland["groups"].items():
            rows.append([f"{gid} — {holland_names.get(gid, gid)}", str(score), "50"])
        self._add_table(doc, rows)
        self._add_bold_paragraph(doc, f"→ Top 3 Holland: {', '.join(holland['top3'])}", size=13)

        # OCEAN table
        self._add_heading2(doc, "TEST 3 — OCEAN / BIG FIVE")
        ocean_names = self._scoring_groups.get("ocean", {})
        rows = [["Chiều tính cách", "Điểm trung bình (1–5)"]]
        for gid, score in ocean["groups"].items():
            rows.append([f"{gid} — {ocean_names.get(gid, gid)}", str(score)])
        self._add_table(doc, rows)

        # SSS table
        self._add_heading2(doc, "SOCIAL SYNCHRONIZATION SCORE (SSS)")
        rows = [
            ["Chỉ số", "Giá trị"],
            ["MBTI Social Score", str(sss["components"]["mbti_social_score"])],
            ["OCEAN E (Hướng ngoại)", str(sss["components"]["ocean_e_avg"])],
            ["Raw Social Score", str(sss["components"]["raw_social_score"])],
            ["SSS (Social Sync Score)", str(sss["score"])],
            ["Social Sync Level", sss["interpretation"]],
        ]
        self._add_table(doc, rows)

    def _add_footer(self, doc: Document, token: str):
        now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        p = doc.add_paragraph()
        run = p.add_run(
            f"Tài liệu nội bộ — Dành cho tư vấn viên · Token: {token} · Ngày tạo: {now_str}"
        )
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ----------------------------------------------------------
    # Formatting helpers
    # ----------------------------------------------------------

    def _group_consecutive_by_axis(self, questions: list) -> List[Tuple[str, list]]:
        """Group questions by axis, preserving order (axes are already contiguous)."""
        groups = []
        current_axis = None
        current = []
        for q in questions:
            if q["axis"] != current_axis:
                if current:
                    groups.append((current_axis, current))
                current_axis = q["axis"]
                current = [q]
            else:
                current.append(q)
        if current:
            groups.append((current_axis, current))
        return groups

    def _add_heading1(self, doc: Document, text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def _add_heading2(self, doc: Document, text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    def _add_field(self, doc: Document, label: str, value: str):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(str(value) if value else "—")

    def _add_rating_answer(self, doc: Document, num: int, question: str, answer: Any):
        p = doc.add_paragraph()
        r1 = p.add_run(f"Câu {num}: ")
        r1.bold = True
        r2 = p.add_run(question)
        r2.italic = True
        r2.font.color.rgb = GRAY
        r3 = p.add_run("\n→ Điểm: ")
        r4 = p.add_run(str(answer))
        r4.bold = True
        r4.font.color.rgb = BLUE

    def _add_bold_paragraph(self, doc: Document, text: str, size: int = 11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)

    def _add_italic_paragraph(self, doc: Document, text: str, color: RGBColor):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = color

    def _add_horizontal_rule(self, doc: Document):
        p = doc.add_paragraph()
        p_fmt = p.paragraph_format
        pPr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        border.append(bottom)
        pPr.append(border)

    def _add_table(self, doc: Document, rows: List[List[str]]):
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r, row_data in enumerate(rows):
            for c, value in enumerate(row_data):
                cell = table.cell(r, c)
                cell.text = str(value)
                for p in cell.paragraphs:
                    for run in p.runs:
                        if r == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
                if r == 0:
                    self._set_cell_background(cell, "1A73E8")
                elif r % 2 == 0:
                    self._set_cell_background(cell, LIGHT_GRAY)

        doc.add_paragraph()  # spacing after table

    def _set_cell_background(self, cell, color_hex: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    def _format_date(self, timestamp_str: str) -> str:
        try:
            dt = date_parser.parse(str(timestamp_str))
            return dt.strftime("%d-%m-%Y")
        except Exception:
            return datetime.now().strftime("%d-%m-%Y")

    # ----------------------------------------------------------
    # Drive API
    # ----------------------------------------------------------

    def _build_drive_service(self):
        creds_json = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
        if not creds_json:
            raise RuntimeError("GOOGLE_SERVICE_ACCOUNT_JSON environment variable not set")
        info  = json.loads(creds_json)
        creds = service_account.Credentials.from_service_account_info(info, scopes=SCOPES)
        return build("drive", "v3", credentials=creds)

    def _get_submission_folder_id(self) -> str:
        parent_id = self._find_or_create_folder(DRIVE_FOLDER_NAME, parent_id=None)
        return self._find_or_create_folder(SUBFOLDER_NAME, parent_id=parent_id)

    def _find_or_create_folder(self, name: str, parent_id: str | None) -> str:
        query = (
            f"name = '{name}' and "
            f"mimeType = 'application/vnd.google-apps.folder' and "
            f"trashed = false"
        )
        if parent_id:
            query += f" and '{parent_id}' in parents"

        results = self._drive.files().list(q=query, fields="files(id, name)").execute()
        files = results.get("files", [])
        if files:
            return files[0]["id"]

        metadata = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
        if parent_id:
            metadata["parents"] = [parent_id]
        folder = self._drive.files().create(body=metadata, fields="id").execute()
        return folder["id"]

    def _upload_as_google_doc(self, buf: io.BytesIO, filename: str, folder_id: str) -> str:
        file_metadata = {
            "name": filename,
            "mimeType": "application/vnd.google-apps.document",
            "parents": [folder_id],
        }
        media = MediaIoBaseUpload(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=True,
        )
        file = self._drive.files().create(
            body=file_metadata, media_body=media, fields="id"
        ).execute()
        return file["id"]
